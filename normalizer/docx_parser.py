"""
docx_parser.py - Parse Word (.docx) security questionnaires into normalized form.

Design decisions:
- python-docx exposes paragraphs and tables separately. We process both.
- Paragraphs are simpler: each paragraph is a potential question or header.
- Tables are trickier: questions are often in col 0, answers in col 1+.
- We track source_location carefully so you can trace back to the original doc.
"""

from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re

from docx import Document

from .detector import detect_response_type, is_question, is_section_header, is_heading_not_question, resolve_section_name
from .utils import clean_text, is_blank, is_noise_row, make_question_id, make_section_id


def parse_docx(filepath: str) -> Dict:
    """
    Entry point: parse a docx file and return normalized structure.
    """
    doc = Document(filepath)

    sections: List[Dict] = []
    current_section: Optional[Dict] = None
    section_counter = 0
    question_counter = 0
    version = ""
    questionnaire_name = ""

    # We iterate over the document body in order, which interleaves paragraphs and tables
    # python-docx doesn't do this natively, so we rebuild the order ourselves.
    body_elements = _iter_body_elements(doc)

    for element_type, element in body_elements:

        if element_type == "paragraph":
            text = clean_text(element.text)

            if is_blank(text):
                continue

            # Grab metadata from early paragraphs
            if not version:
                v = _try_extract_version(text)
                if v:
                    version = v
            if not questionnaire_name:
                n = _try_extract_questionnaire_name(text)
                if n:
                    questionnaire_name = n

            if is_noise_row(text):
                continue

            if is_section_header(text) and (not is_question(text) or is_heading_not_question(text)):
                section_counter += 1
                question_counter = 0
                sec_id = _extract_section_id(text) or make_section_id(section_counter)
                sec_title = resolve_section_name(sec_id)
                current_section = {
                    "section_id": sec_id,
                    "section_title": sec_title,
                    "questions": [],
                }
                sections.append(current_section)
                continue

            if is_question(text):
                current_section, section_counter = _ensure_section(
                    current_section, section_counter, sections, "General"
                )
                question_counter += 1
                q_id = make_question_id(current_section["section_id"], question_counter)
                response_info = detect_response_type(text, [])
                source_loc = {
                    "type": "paragraph",
                    "paragraph_index": element._element.getparent().index(element._element),
                }
                current_section["questions"].append({
                    "question_id": q_id,
                    "text": text,
                    "response_type": response_info["response_type"],
                    "options": response_info["options"],
                    "guidance": "",
                    "source_location": source_loc,
                })

        elif element_type == "table":
            table, table_idx = element
            table_sections, current_section, section_counter, question_counter = _parse_table(
                table, table_idx, current_section, section_counter, question_counter, sections
            )
            sections.extend(table_sections)

    sections = [s for s in sections if s["questions"]]
    return {
        "questionnaire_name": questionnaire_name or Path(filepath).stem,
        "version": version,
        "sections": sections,
    }


def _iter_body_elements(doc):
    """
    Yield (type, element) tuples in document order.
    python-docx splits paragraphs and tables; we re-merge them by walking the XML.
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    table_idx = 0
    para_elements = {p._element: p for p in doc.paragraphs}
    table_elements = {t._element: (t, i) for i, t in enumerate(doc.tables)}

    for child in body:
        if child in para_elements:
            yield "paragraph", para_elements[child]
        elif child in table_elements:
            yield "table", table_elements[child]
            table_idx += 1


def _parse_table(
    table, table_idx: int, current_section, section_counter: int,
    question_counter: int, existing_sections: List[Dict]
) -> Tuple[List[Dict], Optional[Dict], int, int]:
    """
    Parse a single docx table. Returns (new_sections, current_section, section_counter, question_counter).

    Convention: col 0 usually holds question text; col 1+ may hold options or answers.
    """
    new_sections: List[Dict] = []

    for row_idx, row in enumerate(table.rows):
        cells = [clean_text(cell.text) for cell in row.cells]

        if not any(cells):
            continue

        # The first non-empty cell is the candidate question/header
        primary = next((c for c in cells if c), "")
        nearby = [c for c in cells[1:] if c]

        if is_blank(primary) or is_noise_row(primary):
            continue

        if is_section_header(primary) and (not is_question(primary) or is_heading_not_question(primary)):
            section_counter += 1
            question_counter = 0
            sec_id = _extract_section_id(primary) or make_section_id(section_counter)
            sec_title = resolve_section_name(sec_id)
            current_section = {
                "section_id": sec_id,
                "section_title": sec_title,
                "questions": [],
            }
            new_sections.append(current_section)
            continue

        if is_question(primary):
            current_section, section_counter = _ensure_section(
                current_section, section_counter, existing_sections + new_sections, "General"
            )
            question_counter += 1
            q_id = make_question_id(current_section["section_id"], question_counter)
            response_info = detect_response_type(primary, nearby)

            current_section["questions"].append({
                "question_id": q_id,
                "text": primary,
                "response_type": response_info["response_type"],
                "options": response_info["options"],
                "guidance": _extract_guidance_from_cells(cells, primary),
                "source_location": {
                    "type": "table",
                    "table": table_idx,
                    "row": row_idx,
                },
            })

    return new_sections, current_section, section_counter, question_counter


def _ensure_section(
    current_section, section_counter: int, sections: List[Dict], default_title: str
) -> Tuple[Dict, int]:
    """Create an implicit section if we haven't seen a header yet."""
    if current_section is None:
        section_counter += 1
        current_section = {
            "section_id": make_section_id(section_counter),
            "section_title": default_title,
            "questions": [],
        }
        sections.append(current_section)
    return current_section, section_counter


def _try_extract_version(text: str) -> str:
    m = re.search(r"\bv(?:ersion)?\s*(\d+[\.\d]*)", text, re.IGNORECASE)
    return m.group(1) if m else ""


def _try_extract_questionnaire_name(text: str) -> str:
    stripped = text.strip()
    if 5 < len(stripped) < 100 and "?" not in stripped:
        if any(kw in stripped.upper() for kw in ["QUESTIONNAIRE", "ASSESSMENT", "CAIQ", "HECVAT", "NIST", "SOC"]):
            return stripped
    return ""


def _extract_section_id(text: str) -> str:
    m = re.match(r"^([A-Z]{2,8}(-\d+)?)\b", text.strip())
    return m.group(1) if m else ""


def _extract_guidance_from_cells(cells: List[str], question_text: str) -> str:
    for cell in cells:
        if cell and cell != question_text and len(cell) > 80:
            return cell
    return ""
