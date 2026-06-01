"""
create_sample_docx.py - Generate a sample security questionnaire in .docx format.
Run once: python create_sample_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample():
    doc = Document()

    # Title
    title = doc.add_heading("Vendor Security Assessment Questionnaire", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Version 1.3 | Confidential")
    doc.add_paragraph(
        "Instructions: Complete all sections below. "
        "If a control is not applicable to your environment, mark it N/A and briefly explain."
    )

    doc.add_paragraph("")  # spacer

    # ---- Section 1: Network Security (paragraphs) ----
    doc.add_heading("1. Network Security", level=1)

    doc.add_paragraph("1.1 Is your network segmented using firewalls or VLANs?")
    doc.add_paragraph("1.2 Do you perform regular vulnerability scans of your network perimeter?")
    doc.add_paragraph("1.3 Are intrusion detection/prevention systems (IDS/IPS) deployed?")
    doc.add_paragraph("1.4 Describe your patch management process for network devices.")
    doc.add_paragraph("1.5 Are all remote access connections made via VPN with MFA?")

    doc.add_paragraph("")  # spacer

    # ---- Section 2: Data Protection (table) ----
    doc.add_heading("2. Data Protection", level=1)
    doc.add_paragraph("The following questions cover how customer data is handled.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Response"
    hdr[2].text = "Evidence / Notes"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    data_questions = [
        ("Does your organization classify data by sensitivity level?", "Yes / No / N/A", ""),
        ("Is personal data (PII) stored in encrypted form?", "Yes / No / N/A", ""),
        ("Describe your data retention and deletion policy.", "Free text", ""),
        ("Do you have a data loss prevention (DLP) solution in place?", "Yes / No / N/A", ""),
        ("Are data transfers to third parties governed by data processing agreements (DPAs)?", "Yes / No / N/A", ""),
    ]

    for q, r, e in data_questions:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = r
        row[2].text = e

    doc.add_paragraph("")

    # ---- Section 3: Physical Security (mixed) ----
    doc.add_heading("3. Physical Security", level=1)

    doc.add_paragraph("3.1 Are your data centers protected by physical access controls (badge, biometric)?")
    doc.add_paragraph("3.2 Is CCTV surveillance in place for all server room areas?")
    doc.add_paragraph("3.3 Provide details of your visitor management policy for data center access.")

    doc.save("data/sample_questionnaire.docx")
    print("Created data/sample_questionnaire.docx")


if __name__ == "__main__":
    create_sample()
